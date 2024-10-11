from docx import Document

def create_invitations(guests_file):
  """
  This function reads guest names from a file and creates a Word document with custom invitations.

  Args:
    guests_file (str): Path to the file containing guest names.
    docx_template (str): Path to the styled Word document template.
  """
  doc = Document()

  with open(guests_file, 'r') as f:
    guests = f.readlines()

  for guest in guests:
    guest_name = guest.strip()  # Remove trailing newline character

    # Add invitation text with placeholder for guest name
    paragraph = doc.add_paragraph(f"It is a pleasure to have the company of {{guest_name}}", style='Normal')
    paragraph = doc.add_paragraph(guest_name, style='Heading 1')
    paragraph = doc.add_paragraph("at 11010 Memory Lane on the Evening of", style='Normal')
    paragraph = doc.add_paragraph("April 1st", style='Normal')
    paragraph = doc.add_paragraph("at 7 o'clock", style='Normal')

    # Add a page break after each guest invitation
    doc.add_page_break()

  # Save the generated Word document with invitations
  doc.save("custom_invitations.docx")

    

# Replace these paths with your actual file paths
guests_file = "./guests.txt"

create_invitations(guests_file)